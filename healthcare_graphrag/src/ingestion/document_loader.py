"""
Document loader supporting 4 source types:
  research_paper  (.pdf, .txt)
  internal_report (.pdf, .docx, .txt)
  news_article    (.txt, .html)
  email_slack     (.txt, .json)

Every loaded document is tagged with source_type and rich metadata
so downstream modules can apply credibility weighting and filtered retrieval.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


# ── Pydantic models ───────────────────────────────────────────────────────────

class DocumentMetadata(BaseModel):
    """Source-type agnostic metadata envelope."""
    doc_id: str
    title: str
    source_type: str          # research_paper | internal_report | news_article | email_slack
    file_path: str
    date: Optional[str] = None
    # research_paper extras
    authors: List[str] = Field(default_factory=list)
    journal: Optional[str] = None
    doi: Optional[str] = None
    abstract: Optional[str] = None
    # internal_report extras
    department: Optional[str] = None
    patient_cohort: Optional[str] = None
    # news_article extras
    publication: Optional[str] = None
    journalist: Optional[str] = None
    # email_slack extras
    sender_role: Optional[str] = None
    thread_id: Optional[str] = None
    participants: List[str] = Field(default_factory=list)


class LoadedDocument(BaseModel):
    metadata: DocumentMetadata
    raw_text: str


# ── Helpers ───────────────────────────────────────────────────────────────────

def _make_doc_id(file_path: str) -> str:
    return hashlib.md5(file_path.encode()).hexdigest()[:16]


def _extract_date(text: str) -> Optional[str]:
    patterns = [
        r"\b(\d{4}-\d{2}-\d{2})\b",
        r"\b(\w+ \d{1,2},\s*\d{4})\b",
        r"\b(\d{1,2}/\d{1,2}/\d{4})\b",
        r"\b(Q[1-4]\s+\d{4})\b",
    ]
    for p in patterns:
        m = re.search(p, text[:2000])
        if m:
            return m.group(1)
    return None


# ── Source-specific loaders ───────────────────────────────────────────────────

class _PDFLoader:
    @staticmethod
    def load(path: Path) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            logger.warning("pypdf not installed — reading PDF as binary fallback")
            return path.read_text(errors="replace")
        except Exception as exc:
            logger.error("PDF load failed %s: %s", path, exc)
            return ""


class _DocxLoader:
    @staticmethod
    def load(path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not installed — skipping %s", path)
            return ""
        except Exception as exc:
            logger.error("DOCX load failed %s: %s", path, exc)
            return ""


class _TextLoader:
    @staticmethod
    def load(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")


class _HTMLLoader:
    @staticmethod
    def load(path: Path) -> str:
        html = path.read_text(encoding="utf-8", errors="replace")
        # Strip tags naively (no external dep required)
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"&[a-z]+;", " ", text)
        return re.sub(r"\s{2,}", "\n", text).strip()


class _JSONSlackLoader:
    @staticmethod
    def load(path: Path) -> str:
        try:
            data = json.loads(path.read_text())
            if isinstance(data, list):
                lines = []
                for msg in data:
                    user = msg.get("user", msg.get("username", "unknown"))
                    text = msg.get("text", "")
                    ts = msg.get("ts", "")
                    lines.append(f"[{ts}] {user}: {text}")
                return "\n".join(lines)
            return json.dumps(data, indent=2)
        except Exception as exc:
            logger.error("JSON load failed %s: %s", path, exc)
            return ""


# ── Per-source metadata extractors ───────────────────────────────────────────

def _extract_research_paper_meta(text: str, path: Path, doc_id: str) -> DocumentMetadata:
    title = path.stem.replace("_", " ").title()
    authors: List[str] = []
    journal: Optional[str] = None
    doi: Optional[str] = None
    abstract: Optional[str] = None

    # Abstract
    abs_m = re.search(r"(?i)abstract[:\s]+(.{50,1000}?)(?:\n\n|\bintroduction\b)", text)
    if abs_m:
        abstract = abs_m.group(1).strip()

    # DOI
    doi_m = re.search(r"10\.\d{4,9}/\S+", text)
    if doi_m:
        doi = doi_m.group(0).rstrip(".,")

    # Authors — heuristic: line containing comma-separated capitalized names near top
    for line in text.splitlines()[:30]:
        line = line.strip()
        if re.match(r"^([A-Z][a-z]+ [A-Z][a-z]+)(,\s*[A-Z][a-z]+ [A-Z][a-z]+)+", line):
            authors = [a.strip() for a in line.split(",")]
            break

    # Journal
    for kw in ["Journal of", "Annals of", "NEJM", "Lancet", "BMJ", "JAMA", "Nature"]:
        if kw in text[:3000]:
            j_m = re.search(rf"({re.escape(kw)}[^\n]{{0,60}})", text[:3000])
            if j_m:
                journal = j_m.group(1).strip()
                break

    # Title override: first non-empty line
    for line in text.splitlines():
        if len(line.strip()) > 10:
            title = line.strip()[:200]
            break

    return DocumentMetadata(
        doc_id=doc_id,
        title=title,
        source_type="research_paper",
        file_path=str(path),
        date=_extract_date(text),
        authors=authors,
        journal=journal,
        doi=doi,
        abstract=abstract,
    )


def _extract_internal_report_meta(text: str, path: Path, doc_id: str) -> DocumentMetadata:
    title = path.stem.replace("_", " ").title()
    department: Optional[str] = None
    patient_cohort: Optional[str] = None

    dept_m = re.search(r"(?i)department[:\s]+([^\n]{5,60})", text[:2000])
    if dept_m:
        department = dept_m.group(1).strip()

    cohort_m = re.search(r"(?i)(patient cohort|cohort)[:\s]+([^\n]{5,80})", text[:3000])
    if cohort_m:
        patient_cohort = cohort_m.group(2).strip()

    for line in text.splitlines():
        if len(line.strip()) > 10:
            title = line.strip()[:200]
            break

    return DocumentMetadata(
        doc_id=doc_id,
        title=title,
        source_type="internal_report",
        file_path=str(path),
        date=_extract_date(text),
        department=department,
        patient_cohort=patient_cohort,
    )


def _extract_news_article_meta(text: str, path: Path, doc_id: str) -> DocumentMetadata:
    title = path.stem.replace("_", " ").title()
    publication: Optional[str] = None
    journalist: Optional[str] = None

    pub_m = re.search(r"(?i)(published by|source|publication)[:\s]+([^\n]{3,60})", text[:2000])
    if pub_m:
        publication = pub_m.group(2).strip()

    jour_m = re.search(r"(?i)(by|author|reporter)[:\s]+([A-Z][a-z]+ [A-Z][a-z]+)", text[:2000])
    if jour_m:
        journalist = jour_m.group(2).strip()

    for line in text.splitlines():
        if len(line.strip()) > 10:
            title = line.strip()[:200]
            break

    return DocumentMetadata(
        doc_id=doc_id,
        title=title,
        source_type="news_article",
        file_path=str(path),
        date=_extract_date(text),
        publication=publication,
        journalist=journalist,
    )


def _extract_email_slack_meta(text: str, path: Path, doc_id: str) -> DocumentMetadata:
    title = path.stem.replace("_", " ").title()
    sender_role: Optional[str] = None
    thread_id: Optional[str] = None
    participants: List[str] = []

    role_m = re.search(r"(?i)(dr\.|doctor|researcher|nurse|physician)\s+([A-Z][a-z]+)", text[:2000])
    if role_m:
        sender_role = role_m.group(1).strip()

    thread_m = re.search(r"(?i)thread[_\s]?id[:\s]+([A-Za-z0-9_-]+)", text[:2000])
    if thread_m:
        thread_id = thread_m.group(1)

    participants = list(set(re.findall(r"(?:Dr\.|Prof\.)\s+[A-Z][a-z]+", text[:3000])))

    return DocumentMetadata(
        doc_id=doc_id,
        title=title,
        source_type="email_slack",
        file_path=str(path),
        date=_extract_date(text),
        sender_role=sender_role,
        thread_id=thread_id,
        participants=participants,
    )


# ── Source-type detector ──────────────────────────────────────────────────────

_SOURCE_DIR_MAP = {
    "research_papers": "research_paper",
    "internal_reports": "internal_report",
    "news_articles": "news_article",
    "emails_slack": "email_slack",
}


def _detect_source_type(path: Path) -> str:
    for part in path.parts:
        if part in _SOURCE_DIR_MAP:
            return _SOURCE_DIR_MAP[part]
    # Fallback: filename heuristics
    name = path.stem.lower()
    if "report" in name:
        return "internal_report"
    if "news" in name or "article" in name:
        return "news_article"
    if "email" in name or "slack" in name:
        return "email_slack"
    return "research_paper"


# ── Raw text loader ───────────────────────────────────────────────────────────

def _load_raw_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _PDFLoader.load(path)
    if suffix == ".docx":
        return _DocxLoader.load(path)
    if suffix == ".html":
        return _HTMLLoader.load(path)
    if suffix == ".json":
        return _JSONSlackLoader.load(path)
    return _TextLoader.load(path)


# ── Public API ────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".json"}


def load_document(file_path: str | Path) -> LoadedDocument:
    """Load a single document and return a ``LoadedDocument``."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    raw_text = _load_raw_text(path)
    doc_id = _make_doc_id(str(path.resolve()))
    source_type = _detect_source_type(path)

    extractor_map = {
        "research_paper": _extract_research_paper_meta,
        "internal_report": _extract_internal_report_meta,
        "news_article": _extract_news_article_meta,
        "email_slack": _extract_email_slack_meta,
    }
    metadata = extractor_map[source_type](raw_text, path, doc_id)
    return LoadedDocument(metadata=metadata, raw_text=raw_text)


def load_directory(data_dir: str | Path) -> List[LoadedDocument]:
    """Recursively load all supported documents under *data_dir*."""
    data_dir = Path(data_dir)
    documents: List[LoadedDocument] = []

    for path in sorted(data_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                doc = load_document(path)
                documents.append(doc)
                logger.info("Loaded %s (%s)", path.name, doc.metadata.source_type)
            except Exception as exc:
                logger.error("Failed to load %s: %s", path, exc)

    logger.info("Total documents loaded: %d", len(documents))
    return documents
